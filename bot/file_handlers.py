import re
from io import BytesIO

from docx import Document

from bot.consts import TIME_UNITS, TEXT, DURATION


def input_file(file_name):
    """
        Reads a text file containing questions
        and durations, and returns a dictionary.

        The input file should follow the
        specified pattern, for example:
        1. How do you define system requirements based
        on the most successful ways in your experience?
        What methods and tools do you use for this? 2 min 15 sec
        ...

        Each question should be numbered and
        followed by its text and duration in minutes and/or seconds.
    """
    # Read the content of the file
    questions_data = {}
    with open(file_name, "r", encoding="utf-8") as file:
        text = file.read()

    pattern = re.compile(r'(.*?)–?\s+(\d+)\s*(min|sec)(\s+(\d+)\s*sec)?',
                         re.DOTALL | re.UNICODE)
    matches = pattern.finditer(text)

    for question_number, match in enumerate(matches, start=1):
        (question_text, question_duration, time_unit, seconds_part,
         seconds_duration) = match.groups()

        if time_unit not in TIME_UNITS:
            raise ValueError(
                "Invalid time format. "
                "Please enter time in minutes or seconds.")

        duration = int(question_duration) * TIME_UNITS[time_unit]
        if seconds_part:
            duration += int(seconds_duration)

        questions_data[question_number] = {TEXT: question_text,
                                           DURATION: duration}

    return questions_data


def generate_word_document(user_info, answers_data, topic_name):
    word_file = BytesIO()
    doc = Document()

    # Add a heading for the report
    doc.add_heading(
        f"Answers to questions from {user_info['name']} on the topic '{topic_name}'",
        level=1)

    # Add user information to the document
    if 'username' in user_info:
        doc.add_paragraph(
            f"Telegram username: {user_info['username']}")

    # Add a heading for answers
    doc.add_heading("Answers on questions", level=2)

    # Loop through questions and answers and add them to the document
    for question, answer in answers_data.items():
        p_question = doc.add_paragraph()
        run_question = p_question.add_run(f"Question: {question}")
        run_question.bold = True

        doc.add_paragraph(f"Answer: {answer}")
        doc.add_paragraph()  # Add an empty paragraph between questions

    doc.save(word_file)
    word_file.seek(0)

    return word_file.read()


def read_addressee(file_name):
    with open(file_name, "r") as file:
        return [line for line in file]
